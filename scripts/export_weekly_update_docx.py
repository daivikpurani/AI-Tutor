#!/usr/bin/env python3
"""Export the weekly research update (security & improvements) to DOCX for Google Docs."""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Install python-docx: pip install python-docx")
    raise

# Content from docs/WEEKLY_RESEARCH_UPDATE_SECURITY_IMPROVEMENTS.md
CONTENT = {
    "title": "Weekly Research Update",
    "sections": [
        ("Last Week's Tasks", [
            "Implemented security hardening for the AI-Tutor backend: prompt-injection guards, rate limiting, and input sanitization.",
            "Added API and config safeguards: request length limits, filename sanitization, and configurable rejection of suspicious prompts.",
            "Consolidated and cleaned documentation: removed redundant status/progress/benchmark docs and aligned README, QUICK_START, and DEVELOPER with the current stack and security model.",
            "Carried over RAG and retrieval improvements from the previous weekend: OpenAI embeddings, cross-encoder reranking, semantic chunking, and migration tooling for re-indexing.",
        ]),
        ("Completed Work", [
            "Prompt security: Introduced prompt_guard module with injection-pattern detection (e.g. \"ignore previous instructions\", \"reveal your prompt\"), user-message sanitization (strip, truncate to 2000 chars, collapse newlines), and wrapping of user content in <<<USER_QUESTION>>> … <<<END_USER_QUESTION>>> so the model treats it as data. System prompts were updated with explicit \"CRITICAL - Prompt security\" instructions. When REJECT_ON_INJECTION=true, detected injection returns a safe message without calling the LLM. Logging of prompt/response content was stopped to avoid leaking instructions.",
            "API hardening: Chat and upload endpoints are rate-limited via SlowAPI (configurable RATE_LIMIT_CHAT, RATE_LIMIT_UPLOAD in .env). ChatRequest enforces message max length 2000 and conversation_history max 20 entries. All file upload paths use a new filename_sanitizer: strip path components, remove control characters, validate extension against allowed types, cap length at 255. In production, generic 500 responses avoid exposing internals.",
            "Config and tests: Added backend_python/.env.example with REJECT_ON_INJECTION, rate limits, and CORS; all references updated from env.example to .env.example. New tests: test_prompt_guard.py (detection, sanitize, wrap) and test_chat_rejects_injection_when_enabled in test_api.py.",
            "Docs consolidation: README, QUICK_START, and DEVELOPER updated for current stack and security. Removed redundant status/progress/benchmark/startup docs. Single entry point for benchmarking: docs/BENCHMARK.md plus docs/BENCHMARK_DETAILED.md; scripts/benchmarks/README.md for benchmark scripts. PDF generation uses BENCHMARK_DETAILED.",
            "RAG improvements (previous weekend): OpenAI embeddings and cross-encoder reranking for better retrieval; semantic chunking (LangChain) in document chunker; configurable retrieval parameters and thresholds; migration script for re-indexing with new embeddings/chunking; query-handler and vector-db updates documented in implementation and testing notes.",
        ]),
        ("Security & Improvement Areas (What We Implemented)", [
            ("Prompt and input guards", [
                "Injection-pattern detection with configurable reject-on-detect.",
                "User-message sanitization (length, control chars, newlines) and delimiter wrapping in all prompts.",
                "Hardened system prompts that instruct the model to treat only delimited content as the user question.",
            ]),
            ("API and request safety", [
                "SlowAPI rate limits on chat and upload endpoints (e.g. 60/min chat, 10/min upload).",
                "Pydantic request limits: message length, conversation history length.",
                "CORS configuration via .env (e.g. CORS_ORIGINS).",
            ]),
            ("File and filename safety", [
                "Dedicated filename_sanitizer: no path traversal, no control chars, allowed-extension checks, max filename length.",
                "Applied to all upload flows (single, multiple, and title-based filename creation).",
            ]),
            ("Operations and documentation", [
                "Production mode: generic 500, no prompt/response logging.",
                "Single .env.example with security and CORS; DEVELOPER and README document rate limits, prompt security, and optional rejection.",
            ]),
        ]),
        ("Summary of Security & Hardening", [
            "The work is organized into four layers:",
            ("Prompt layer", [
                "Wraps user input in delimiters and checks for injection patterns.",
                "Optional rejection with a safe message when injection is detected.",
                "Reduces risk of instruction override and prompt leaking.",
            ]),
            ("API layer", [
                "Rate limits on chat and upload to curb abuse and cost.",
                "Request schema limits (message length, history size) to keep payloads bounded.",
                "CORS restricted to configured origins.",
            ]),
            ("Input layer", [
                "Sanitized user messages (length, control chars) before inclusion in prompts.",
                "Sanitized filenames (path stripping, extension allowlist) on all uploads.",
                "Ensures only safe, validated input reaches the RAG pipeline and storage.",
            ]),
            ("Docs and ops layer", [
                "Consolidated benchmark and developer docs; single BENCHMARK entry + BENCHMARK_DETAILED.",
                "Security and config documented in README, DEVELOPER, and .env.example.",
                "Production guidance: ENVIRONMENT=production, strong SECRET_KEY, no sensitive logging.",
            ]),
        ]),
        ("Insights", [
            "Prompt injection detection plus delimiter wrapping significantly narrows the attack surface for instruction override and prompt extraction; optional reject mode gives a strict option for higher-risk deployments.",
            "Rate limiting and request limits are simple to configure (SlowAPI + Pydantic) and protect both cost and availability without changing core RAG logic.",
            "Filename sanitization and extension checks prevent path traversal and arbitrary file types; one shared module keeps behavior consistent across all upload endpoints.",
            "Documentation consolidation (single benchmark entry, one DEVELOPER guide, aligned README) makes it easier for new contributors and evaluators to find security and run instructions.",
        ]),
        ("Next Week", [
            "Continue reading and refining the benchmarking methodology (BENCHMARK_DETAILED, evaluation layers, and scripts).",
            "Run a small-scale benchmark pass with the current security and RAG stack to confirm no regressions and document any latency/behavior notes.",
        ]),
        ("Future Updates (Planned)", [
            "User authentication and sessions — Add optional auth (e.g. API key or session-based) and role-based access so that rate limits and logging can be scoped per user and prepared for multi-tenant or lab deployments.",
            "Stricter production defaults — Consider defaulting REJECT_ON_INJECTION=true and tighter rate limits when ENVIRONMENT=production, with clear docs and migration steps for existing deployments.",
        ]),
    ],
}


def build_doc():
    doc = Document()
    doc.add_heading(CONTENT["title"], 0)

    for section_title, section_items in CONTENT["sections"]:
        doc.add_heading(section_title, level=1)
        for item in section_items:
            if isinstance(item, str):
                doc.add_paragraph(item, style="List Bullet")
            else:
                subheading, subitems = item
                p = doc.add_paragraph()
                run = p.add_run(subheading)
                run.bold = True
                for sub in subitems:
                    doc.add_paragraph(sub, style="List Bullet")

    return doc


def main():
    root = Path(__file__).resolve().parent.parent
    out_path = root / "docs" / "WEEKLY_RESEARCH_UPDATE_SECURITY_IMPROVEMENTS.docx"
    doc = build_doc()
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
